"""
Raw text extraction for uploaded academic documents.

Kept deliberately small and dependency-light: PDF and DOCX are the two
formats students realistically upload for timetables/syllabi/calendars;
plain text covers pasted notes. Extend here (not by touching the service
layer) if more formats are needed later.
"""
from app.core.exceptions import CustomException


def extract_text(file_bytes: bytes, content_type: str, filename: str) -> str:
    if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.lower().endswith(".docx")
    ):
        return _extract_docx(file_bytes)
    if content_type == "text/plain" or filename.lower().endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")

    raise CustomException(
        f"Unsupported document type '{content_type}'. Please upload a PDF, DOCX, or TXT file.",
        status_code=400,
    )


def _extract_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    from io import BytesIO

    reader = PdfReader(BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_bytes: bytes) -> str:
    import docx
    from io import BytesIO

    document = docx.Document(BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
